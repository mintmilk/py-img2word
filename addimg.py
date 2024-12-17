import os
import zipfile
from docx import Document
from docx.shared import Inches
import imghdr

def extract_images_to_word(zip_folder_path, output_word_path):
    # 创建Word文档
    doc = Document()
    
    # 遍历文件夹中的所有文件
    for filename in os.listdir(zip_folder_path):
        if filename.endswith(('.zip', '.7z', '.rar')):
            zip_path = os.path.join(zip_folder_path, filename)
            
            try:
                # 创建临时解压目录
                temp_extract_path = os.path.join(zip_folder_path, 'temp_extract')
                os.makedirs(temp_extract_path, exist_ok=True)
                
                # 解压文件
                with zipfile.ZipFile(zip_path, 'r') as zip_ref:
                    zip_ref.extractall(temp_extract_path)
                
                # 获取压缩包前14个字符作为标题
                title = filename[:14]
                
                # 添加标题到Word文档
                doc.add_heading(title, level=2)
                
                # 遍历解压目录中的图片
                image_files = [f for f in os.listdir(temp_extract_path) 
                               if imghdr.what(os.path.join(temp_extract_path, f)) is not None]
                
                # 将图片添加到Word文档
                for img_name in image_files:
                    img_path = os.path.join(temp_extract_path, img_name)
                    doc.add_picture(img_path, width=Inches(6))
                
                # 清理临时目录
                for item in os.listdir(temp_extract_path):
                    os.remove(os.path.join(temp_extract_path, item))
                os.rmdir(temp_extract_path)
            
            except Exception as e:
                print(f"处理 {filename} 时出错: {e}")
    
    # 保存Word文档
    doc.save(output_word_path)
    print(f"已将图片保存到 {output_word_path}")

# 使用示例
zip_folder = './02'
output_word = 'document2.docx'
extract_images_to_word(zip_folder, output_word)